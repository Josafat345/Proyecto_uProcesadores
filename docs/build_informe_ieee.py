from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


PROJECT_DIR = Path(r"C:\Users\josaf\OneDrive\Documents\VII Semestre\05_uProcesadores\07_Proyecto")
DOCS_DIR = PROJECT_DIR / "docs"
FIG_DIR = DOCS_DIR / "figures"
DRAWIO_DIR = DOCS_DIR / "drawio"
OUT_PATH = DOCS_DIR / "Informe_Proyecto_IEEE_JosafatVasquez.docx"
CODE_IMAGE_DIR = DOCS_DIR / "code_appendix_pages"

VERILOG_FILES = [
    PROJECT_DIR / "rtl" / "register4.v",
    PROJECT_DIR / "rtl" / "mux4.v",
    PROJECT_DIR / "rtl" / "alu4.v",
    PROJECT_DIR / "rtl" / "flag_registers.v",
    PROJECT_DIR / "rtl" / "datapath.v",
    PROJECT_DIR / "rtl" / "microstore.v",
    PROJECT_DIR / "rtl" / "control_unit.v",
    PROJECT_DIR / "rtl" / "proyecto_microcodificado_top.v",
    PROJECT_DIR / "tb" / "tb_proyecto_microcodificado.v",
]

BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(70, 70, 70)
LIGHT_GRAY = "EFEFEF"


def set_run_font(run, name="Times New Roman", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_columns(section, count=1, space_twips=360):
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    if count == 1:
        cols.attrib.pop(qn("w:num"), None)
        cols.set(qn("w:space"), str(space_twips))
    else:
        cols.set(qn("w:num"), str(count))
        cols.set(qn("w:space"), str(space_twips))


def set_page(section, columns=1):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    set_columns(section, columns, 360 if columns == 2 else 720)


def shade_cell(cell, fill=LIGHT_GRAY):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=45, start=60, bottom=45, end=60):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    for row in table.rows:
        for i, width in enumerate(widths_in):
            cell = row.cells[i]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r, size=9)
    return p


def add_noindent(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=9, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=9)
    else:
        r = p.add_run(text)
        set_run_font(r, size=9)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    set_run_font(r, size=10, bold=True)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_run_font(r, size=9, bold=True, italic=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r, size=8, italic=True, color=GRAY)
    return p


def add_figure(doc, path, caption, width):
    path = Path(path)
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_table(doc, headers, rows, widths, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=font_size, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == len(row) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            set_run_font(r, size=font_size)
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_code_listing(doc, path):
    add_subheading(doc, path.name)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run()
    set_run_font(r, name="Consolas", size=4.7)
    for line in path.read_text(encoding="utf-8").expandtabs(4).splitlines():
        r.add_text(line if line else " ")
        r.add_break()


def render_code_file_pages(path):
    CODE_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    lines = path.read_text(encoding="utf-8").expandtabs(4).splitlines()
    if not lines:
        lines = [""]

    font_path = Path(r"C:\Windows\Fonts\consola.ttf")
    try:
        font = ImageFont.truetype(str(font_path), 14)
        title_font = ImageFont.truetype(str(font_path), 18)
    except OSError:
        font = ImageFont.load_default()
        title_font = ImageFont.load_default()

    page_width = 1500
    margin_x = 38
    margin_y = 32
    line_height = 16
    title_height = 30
    lines_per_page = 95
    outputs = []

    for page_idx, start in enumerate(range(0, len(lines), lines_per_page), start=1):
        chunk = lines[start:start + lines_per_page]
        page_height = margin_y * 2 + title_height + len(chunk) * line_height
        img = Image.new("RGB", (page_width, page_height), "white")
        draw = ImageDraw.Draw(img)
        draw.rectangle((0, 0, page_width - 1, page_height - 1), outline=(210, 210, 210), width=2)
        title = path.name if len(outputs) == 0 else f"{path.name} (continuación {page_idx})"
        draw.text((margin_x, margin_y), title, fill=(0, 0, 0), font=title_font)
        y = margin_y + title_height
        for line in chunk:
            draw.text((margin_x, y), line if line else " ", fill=(0, 0, 0), font=font)
            y += line_height
        out = CODE_IMAGE_DIR / f"{path.stem}_page_{page_idx}.jpg"
        img.save(out, quality=90, optimize=True)
        outputs.append(out)
    return outputs


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(3)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.color.rgb = BLACK


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Diseño y Simulación de una Trayectoria de Datos Microcodificada de 4 bits")
    set_run_font(r, size=16, bold=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(
        "Josafat Vásquez (9-764-2024)\n"
        "Facultad de Ingeniería Eléctrica\n"
        "Universidad Tecnológica de Panamá\n"
        "Ciudad de Panamá, Panamá\n"
        "josafat.vasquez@utp.com.pa"
    )
    set_run_font(r, size=10)

    add_noindent(
        doc,
        "Resumen— En este proyecto se diseñó y simuló una trayectoria de datos microcodificada de 4 bits en Verilog. "
        "El sistema integra registros, un multiplexor, una ALU, registros de banderas y una unidad de control basada en microcódigo. "
        "La simulación en ModelSim permitió verificar el avance de la dirección de microprograma, la carga selectiva de registros y las bifurcaciones por banderas.",
        "Resumen—",
    )
    add_noindent(
        doc,
        "Palabras clave— microcódigo, datapath, Verilog, ModelSim, ALU, unidad de control, banderas, microstore.",
        "Palabras clave—",
    )


def build_doc():
    doc = Document()
    setup_styles(doc)
    set_page(doc.sections[0], columns=1)
    add_title_block(doc)

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_page(body_section, columns=2)

    add_heading(doc, "I. Introducción")
    add_body(
        doc,
        "Una trayectoria de datos es el conjunto de registros, buses y bloques combinacionales que permiten mover y transformar información dentro de un procesador. "
        "En un sistema microcodificado, las señales que controlan dicha trayectoria no se generan directamente desde una máquina de estados extensa, sino desde una memoria de control que almacena microinstrucciones."
    )
    add_body(
        doc,
        "El objetivo de este proyecto fue implementar un procesador didáctico de 4 bits capaz de ejecutar un pequeño programa de demostración. "
        "El diseño se separó en módulos para facilitar su verificación: registros, multiplexor, ALU, banderas, datapath completo, microstore, unidad de control y banco de pruebas."
    )

    add_heading(doc, "II. Trayectoria de datos")
    add_body(
        doc,
        "La trayectoria de datos contiene tres registros principales, identificados como A, B y C. El registro C alimenta la entrada X de la ALU, mientras que la entrada Y se obtiene desde un multiplexor de cuatro entradas. "
        "El resultado de la ALU, denominado Z, regresa a los registros A, B y C; sin embargo, cada registro se actualiza únicamente cuando su señal de habilitación está activa."
    )
    add_figure(
        doc,
        DRAWIO_DIR / "datapath_profesor_style_full_flags.drawio.png",
        "Figura 1. Trayectoria de datos microcodificada de 4 bits.",
        3.25,
    )
    add_table(
        doc,
        ["SelMux", "Entrada Y"],
        [("00", "A"), ("01", "B"), ("10", "KVAL"), ("11", "SW")],
        [0.7, 2.55],
    )
    add_table(
        doc,
        ["SelAlu", "Operación"],
        [("00", "Z = X + Y"), ("01", "Z = X - Y"), ("10", "Z = Y"), ("11", "Z = X AND Y")],
        [0.7, 2.55],
    )
    add_body(
        doc,
        "Las banderas C y Z se generan a partir de la salida de la ALU. La bandera Z indica resultado cero, mientras que la bandera C se usa como acarreo en suma y como indicador de préstamo en resta. "
        "La señal FS permite seleccionar si las banderas visibles corresponden a la operación actual o al valor previamente almacenado."
    )

    add_heading(doc, "III. Unidad de control")
    add_body(
        doc,
        "La unidad de control es microprogramada. Su contador de dirección produce ADDx, que selecciona una palabra dentro del microstore. "
        "Esa palabra contiene señales para el datapath y campos de secuencia que determinan si la ejecución continúa con la siguiente dirección o si se carga una dirección de salto."
    )
    add_figure(doc, FIG_DIR / "control_unit.png", "Figura 2. Unidad de control microprogramada.", 3.25)
    add_table(
        doc,
        ["Campo", "Descripción"],
        [
            ("ADDx", "Dirección actual del microprograma."),
            ("TEST", "Selecciona la condición de bifurcación."),
            ("NATT", "Dirección destino si se toma el salto."),
            ("EnA/EnB/EnC", "Habilitan carga de registros."),
            ("SelMux/SelAlu", "Seleccionan fuente de datos y operación."),
            ("FS", "Controla la selección de banderas."),
        ],
        [0.95, 2.3],
    )
    add_table(
        doc,
        ["TEST", "Condición"],
        [("00", "No salta; ADDx+1."), ("01", "Salto incondicional."), ("10", "Salta si C-flag=1."), ("11", "Salta si Z-flag=1.")],
        [0.7, 2.55],
    )

    add_heading(doc, "IV. Microcódigo simbólico")
    add_body(
        doc,
        "El microcódigo simbólico describe las operaciones de alto nivel que se almacenan posteriormente como campos binarios dentro del microstore. "
        "Cada línea corresponde a una dirección de control y se traduce en enables, selectores de MUX, operación de ALU, constante KVAL y condición de salto."
    )
    add_subheading(doc, "Lista simbólica de demostración")
    for line in [
        "0000 Start: C <- F",
        "0001 GetSW: A, B <- SW",
        "0010        B, C <- C - A; if Z-flag* goto Start",
        "0011        C <- A",
        "0100 Top:   C <- C + 1",
        "0101        C - B; if C-flag goto Top",
        "0110        goto Start",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=7.2)
    add_body(
        doc,
        "La marca Z-flag* indica que la decisión utiliza una bandera almacenada previamente. En la demostración, el testbench inicia con SW=0000 para comprobar el retorno a Start y luego cambia a SW=0100 para observar el recorrido normal del programa."
    )
    add_figure(doc, FIG_DIR / "asm_chart_entrada.png", "Figura 3. Carta ASM: lectura y decisión inicial.", 3.25)
    add_figure(doc, FIG_DIR / "asm_chart_conteo.png", "Figura 4. Carta ASM: lazo de conteo.", 3.25)

    add_heading(doc, "V. Contenido del Microstore")
    add_body(
        doc,
        "El microstore implementa una ROM combinacional indexada por ADDx. Las direcciones no utilizadas se redirigen al inicio para mantener el sistema en una trayectoria de ejecución definida."
    )
    add_table(
        doc,
        ["Dir", "TEST", "NATT", "En", "SM", "SA", "FS", "KVAL", "Operación"],
        [
            ("0000", "00", "0000", "001", "10", "10", "0", "1111", "C <- F"),
            ("0001", "00", "0000", "110", "11", "10", "1", "0000", "A,B <- SW"),
            ("0010", "11", "0000", "011", "00", "01", "0", "0000", "B,C <- C-A; if Z*"),
            ("0011", "00", "0000", "001", "00", "10", "0", "0000", "C <- A"),
            ("0100", "00", "0000", "001", "10", "00", "0", "0001", "C <- C+1"),
            ("0101", "10", "0100", "000", "01", "01", "1", "0000", "C-B; if C"),
            ("0110", "01", "0000", "000", "00", "00", "0", "0000", "goto Start"),
        ],
        [0.36, 0.36, 0.46, 0.35, 0.35, 0.35, 0.3, 0.45, 0.87],
        font_size=6.4,
    )
    add_noindent(doc, "Convención: En = EnA EnB EnC, SM = SelMux y SA = SelAlu.", "Convención:")

    add_heading(doc, "VI. Tabla de acontecimientos")
    add_body(
        doc,
        "Para hacer la tabla de acontecimientos más legible se agruparon los ciclos por función. Los valores A, B y C representan el contenido observado después del flanco activo del reloj; ADDx representa la dirección de microinstrucción visible durante el ciclo."
    )
    add_table(
        doc,
        ["Ciclo", "ADDx", "Evento", "Resultado observado"],
        [
            ("0-1", "0000", "Reset activo", "A=B=C=0; banderas en 0."),
            ("2", "0001", "Carga inicial", "C=F desde KVAL."),
            ("3", "0010", "SW=0", "Z-flag* provoca retorno a Start."),
            ("6", "0010", "SW=4", "A=4, B=4; continúa ejecución."),
            ("7-8", "0011/0100", "Preparación", "C toma A y luego se incrementa."),
            ("9-21", "0101/0100", "Lazo Top", "C aumenta mientras C-flag=1."),
            ("22-23", "0110/0000", "Fin de vuelta", "Salto incondicional a Start."),
        ],
        [0.5, 0.55, 0.95, 1.25],
        font_size=7,
    )

    add_heading(doc, "VII. Conclusiones generales")
    add_body(
        doc,
        "El proyecto permitió comprobar el funcionamiento de una arquitectura microcodificada sencilla. La separación entre datapath y unidad de control facilita identificar qué bloque procesa datos y qué bloque genera las señales de gobierno."
    )
    add_body(
        doc,
        "La simulación en ModelSim verificó que el microstore produce las señales esperadas, que los registros sólo cargan cuando sus enables están activos y que las bifurcaciones dependen correctamente de TEST, NATT y las banderas C/Z."
    )
    add_body(
        doc,
        "El uso de banderas guardadas mediante FS fue un punto importante del diseño, ya que permitió demostrar la diferencia entre una bandera generada en el ciclo actual y una bandera retenida para una decisión posterior."
    )

    add_heading(doc, "Referencias")
    refs = [
        "[1] S. Palnitkar, Verilog HDL: A Guide to Digital Design and Synthesis, 2nd ed. Prentice Hall, 2003.",
        "[2] Siemens EDA, ModelSim Intel FPGA Edition User Manual.",
        "[3] GOWIN Semiconductor, Tang Nano 9K FPGA Development Board Documentation.",
    ]
    for ref in refs:
        add_noindent(doc, ref)

    appendix_results_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page(appendix_results_section, columns=2)
    add_heading(doc, "Apéndice A. Resultados actuales de simulación")
    add_body(
        doc,
        "Las siguientes figuras muestran la corrida actual de simulación en ModelSim. La consola presenta la tabla de valores por ciclo y la ventana Wave muestra señales de entrada, registros, banderas y control microcodificado."
    )
    add_figure(doc, PROJECT_DIR / "Out.png", "Figura A1. Salida de consola de ModelSim.", 3.25)
    add_figure(doc, PROJECT_DIR / "Wave.png", "Figura A2. Formas de onda de la simulación.", 3.25)

    code_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page(code_section, columns=2)
    add_heading(doc, "Apéndice B. Código fuente Verilog")
    add_body(
        doc,
        "La Tabla B1 lista los módulos Verilog del diseño y el archivo de simulación. En las páginas siguientes se presenta el código fuente RTL actual de forma compacta. Cada archivo RTL incluye un bloque de comentario con nombre del archivo, autor, fecha y descripción breve."
    )
    add_table(
        doc,
        ["Archivo", "Contenido"],
        [
            ("register4.v", "Registro de 4 bits con reset y enable."),
            ("mux4.v", "Multiplexor de entrada Y."),
            ("alu4.v", "ALU de 4 bits."),
            ("flag_registers.v", "Registro/selección de banderas."),
            ("datapath.v", "Trayectoria de datos completa."),
            ("microstore.v", "ROM de microcódigo."),
            ("control_unit.v", "Unidad de control microprogramada."),
            ("proyecto_microcodificado_top.v", "Módulo superior."),
            ("tb_proyecto_microcodificado.v", "Banco de pruebas para ModelSim; disponible en el repositorio."),
        ],
        [1.35, 1.9],
        font_size=7,
    )
    doc.add_section(WD_SECTION.CONTINUOUS)
    set_page(doc.sections[-1], columns=2)
    add_noindent(
        doc,
        "El código fuente se imprime como imágenes separadas por archivo, manteniendo el formato de dos columnas del informe.",
        "El código fuente",
    )
    fig_idx = 1
    for path in VERILOG_FILES:
        for page_idx, image_path in enumerate(render_code_file_pages(path), start=1):
            suffix = "" if page_idx == 1 else f", continuación {page_idx}"
            add_figure(doc, image_path, f"Figura B{fig_idx}. Código fuente: {path.name}{suffix}.", 3.15)
            fig_idx += 1

    constraints_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page(constraints_section, columns=2)
    add_heading(doc, "Apéndice C. Archivo CST/UCF")
    add_body(
        doc,
        "En esta etapa del proyecto no se realizaron cambios a un archivo CST o UCF porque la entrega actual corresponde a la simulación en ModelSim. "
        "Por lo tanto, no existe todavía una asignación física de pines para FPGA dentro del repositorio."
    )
    add_body(
        doc,
        "Cuando el diseño se lleve a la Tang Nano 9K, este apéndice deberá actualizarse con el archivo de restricciones correspondiente, indicando las señales asignadas a reloj, reset, switches, LEDs y cualquier salida de depuración."
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_doc())
